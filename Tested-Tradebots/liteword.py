import tkinter as tk
from tkinter import filedialog, font, messagebox, simpledialog
from tkinter import ttk
from PIL import Image, ImageTk
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

class LiteWord:
    def __init__(self, root):
        self.root = root
        self.root.title("LiteWord - MS Word Clone")
        self.root.geometry("1000x700")
        self.filename = None
        self.images = []
        self.image_paths = []
        self.content_blocks = []  # list of (type, content)
        self.setup_ui()

    def setup_ui(self):
        self.current_font = tk.StringVar(value="Arial")
        self.current_size = tk.IntVar(value=12)

        toolbar = ttk.Frame(self.root)
        toolbar.pack(side=tk.TOP, fill=tk.X)

        font_menu = ttk.Combobox(toolbar, textvariable=self.current_font,
                                 values=font.families(), state='readonly', width=20)
        font_menu.pack(side=tk.LEFT, padx=5)
        font_menu.bind("<<ComboboxSelected>>", self.change_font)

        size_menu = ttk.Combobox(toolbar, textvariable=self.current_size,
                                 values=list(range(8, 72, 2)), width=5)
        size_menu.pack(side=tk.LEFT)
        size_menu.bind("<<ComboboxSelected>>", self.change_font)

        ttk.Button(toolbar, text="Bold", command=self.make_bold).pack(side=tk.LEFT, padx=5)
        ttk.Button(toolbar, text="Italic", command=self.make_italic).pack(side=tk.LEFT)
        ttk.Button(toolbar, text="Underline", command=self.make_underline).pack(side=tk.LEFT)
        ttk.Button(toolbar, text="Bullets", command=self.insert_bullets).pack(side=tk.LEFT)
        ttk.Button(toolbar, text="Numbered", command=self.insert_numbered_list).pack(side=tk.LEFT)
        ttk.Button(toolbar, text="Insert Table", command=self.insert_table).pack(side=tk.LEFT)
        ttk.Button(toolbar, text="Export .docx", command=self.export_to_docx).pack(side=tk.LEFT)

        menubar = tk.Menu(self.root)
        self.root.config(menu=menubar)

        file_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="File", menu=file_menu)
        file_menu.add_command(label="New", command=self.new_file)
        file_menu.add_command(label="Open", command=self.open_file)
        file_menu.add_command(label="Save", command=self.save_file)
        file_menu.add_command(label="Save As", command=self.save_as_file)
        file_menu.add_separator()
        file_menu.add_command(label="Insert Image", command=self.insert_image)
        file_menu.add_separator()
        file_menu.add_command(label="Exit", command=self.root.quit)

        page_frame = ttk.Frame(self.root, padding=(80, 60, 80, 60))
        page_frame.pack(fill=tk.BOTH, expand=1)

        scroll = ttk.Scrollbar(page_frame)
        scroll.pack(side=tk.RIGHT, fill=tk.Y)

        self.text = tk.Text(page_frame, wrap=tk.WORD, undo=True,
                            font=(self.current_font.get(), self.current_size.get()),
                            spacing3=10, padx=40, pady=20, relief=tk.FLAT)
        self.text.pack(side=tk.LEFT, fill=tk.BOTH, expand=1)
        self.text.config(yscrollcommand=scroll.set)
        scroll.config(command=self.text.yview)

    def change_font(self, event=None):
        self.text.configure(font=(self.current_font.get(), self.current_size.get()))

    def make_bold(self):
        self.toggle_tag("bold", {"font": (self.current_font.get(), self.current_size.get(), "bold")})

    def make_italic(self):
        self.toggle_tag("italic", {"font": (self.current_font.get(), self.current_size.get(), "italic")})

    def make_underline(self):
        self.toggle_tag("underline", {"font": (self.current_font.get(), self.current_size.get(), "underline")})

    def toggle_tag(self, tag, config):
        try:
            current_tags = self.text.tag_names("sel.first")
            if tag in current_tags:
                self.text.tag_remove(tag, "sel.first", "sel.last")
            else:
                self.text.tag_configure(tag, **config)
                self.text.tag_add(tag, "sel.first", "sel.last")
        except tk.TclError:
            pass

    def new_file(self):
        self.text.delete(1.0, tk.END)
        self.filename = None
        self.images.clear()
        self.image_paths.clear()
        self.content_blocks.clear()
        self.root.title("LiteWord - New Document")

    def open_file(self):
        path = filedialog.askopenfilename(title="Open File", filetypes=[("Text Files", "*.txt"), ("All Files", "*.*")])
        if path:
            self.text.delete(1.0, tk.END)
            self.images.clear()
            self.image_paths.clear()
            self.content_blocks.clear()
            try:
                with open(path, "r", encoding="utf-8") as f:
                    self.text.insert(tk.END, f.read())
                self.filename = path
                self.root.title(f"LiteWord - {os.path.basename(path)}")
            except Exception as e:
                messagebox.showerror("Open Error", str(e))

    def save_file(self):
        if self.filename:
            self._save_as_txt(self.filename)
        else:
            self.save_as_file()

    def save_as_file(self):
        path = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[("Text Files", "*.txt")])
        if path:
            self._save_as_txt(path)
            self.filename = path
            self.root.title(f"LiteWord - {os.path.basename(path)}")

    def _save_as_txt(self, path):
        try:
            with open(path, "w", encoding="utf-8") as f:
                f.write(self.text.get(1.0, tk.END))
        except Exception as e:
            messagebox.showerror("Save Error", str(e))

    def insert_image(self):
        path = filedialog.askopenfilename(filetypes=[("Image Files", "*.png *.jpg *.jpeg *.gif")])
        if path:
            try:
                img = Image.open(path)
                img.thumbnail((400, 400))
                img_tk = ImageTk.PhotoImage(img)
                self.images.append(img_tk)
                self.image_paths.append(path)
                self.text.image_create(tk.END, image=img_tk)
                self.text.insert(tk.END, "\n")
                self.content_blocks.append(("image", path))
            except Exception as e:
                messagebox.showerror("Image Error", str(e))

    def insert_bullets(self):
        try:
            self.text.insert(tk.INSERT, "\u2022 ")
        except:
            pass

    def insert_numbered_list(self):
        try:
            current = self.text.index(tk.INSERT).split(".")[0]
            self.text.insert(tk.INSERT, f"{int(current)}. ")
        except:
            pass

    def insert_table(self):
        rows = simpledialog.askinteger("Rows", "Enter number of rows:")
        cols = simpledialog.askinteger("Columns", "Enter number of columns:")
        if rows and cols:
            self.content_blocks.append(("table", (rows, cols)))
            self.text.insert(tk.END, f"[Table: {rows}x{cols}]\n")

    def export_to_docx(self):
        path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
        if path:
            doc = Document()
            section = doc.sections[-1]
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)

            style = doc.styles['Normal']
            style.font.name = 'Arial'
            style.font.size = Pt(12)

            lines = self.text.get(1.0, tk.END).split("\n")
            img_index = 0
            table_index = 0
            for line in lines:
                if line.startswith("[Table: "):
                    rows, cols = self.content_blocks[table_index][1]
                    table = doc.add_table(rows=rows, cols=cols)
                    for r in range(rows):
                        for c in range(cols):
                            table.cell(r, c).text = f"Cell {r+1},{c+1}"
                    table_index += 1
                elif "image" in line.lower() and img_index < len(self.image_paths):
                    doc.add_picture(self.image_paths[img_index], width=Inches(4.5))
                    img_index += 1
                elif line.strip():
                    p = doc.add_paragraph(line)
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            doc.save(path)
            messagebox.showinfo("Success", f"Exported to {path}")

if __name__ == "__main__":
    root = tk.Tk()
    app = LiteWord(root)
    root.mainloop()
